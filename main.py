import streamlit as st
from openai import OpenAI
import docx
from io import BytesIO
import re

# Initialize OpenAI client
client = OpenAI()

st.title("Mask AI")

# File uploader for txt or docx files
uploaded_file = st.file_uploader("Upload a txt or docx file", type=["txt", "docx"])

# Function to mask personal and company names using ChatGPT
def mask_with_chatgpt(text):
    prompt = (
        "Please mask all personal names and company names in the following text. "
        "Use [MASKED NAME] for personal names and [MASKED COMPANY] for company names:\n\n"
        f"{text}"
    )

    completion = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ]
    )

    masked_text = completion.choices[0].message.content
    return masked_text

# Function to mask docx file content
def mask_docx(doc, masked_text):
    masked_doc = docx.Document()
    
    original_paragraphs = [para.text for para in doc.paragraphs]
    masked_paragraphs = masked_text.split('\n')

    for orig_para, masked_para in zip(doc.paragraphs, masked_paragraphs):
        new_para = masked_doc.add_paragraph()
        
        # Preserve original paragraph formatting
        new_para.paragraph_format.left_indent = orig_para.paragraph_format.left_indent
        new_para.paragraph_format.right_indent = orig_para.paragraph_format.right_indent
        new_para.paragraph_format.first_line_indent = orig_para.paragraph_format.first_line_indent
        new_para.paragraph_format.alignment = orig_para.paragraph_format.alignment
        
        masked_parts = re.split(r'(\[MASKED NAME\]|\[MASKED COMPANY\])', masked_para)
        orig_cursor = 0

        for part in masked_parts:
            if part in ["[MASKED NAME]", "[MASKED COMPANY]"]:
                # Add masked part as a new run and preserve original formatting
                for run in orig_para.runs:
                    if orig_para.text[orig_cursor:].startswith(run.text):
                        new_run = new_para.add_run(part)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = run.font.size
                        new_run.font.name = run.font.name
                        new_run.font.color.rgb = run.font.color.rgb
                        orig_cursor += len(run.text)
                        break
            else:
                while orig_cursor < len(orig_para.text):
                    for run in orig_para.runs:
                        run_text = run.text
                        if orig_para.text[orig_cursor:].startswith(run_text):
                            if part in run_text:
                                start_index = run_text.find(part)
                                end_index = start_index + len(part)
                                new_run = new_para.add_run(part)
                                new_run.bold = run.bold
                                new_run.italic = run.italic
                                new_run.underline = run.underline
                                new_run.font.size = run.font.size
                                new_run.font.name = run.font.name
                                new_run.font.color.rgb = run.font.color.rgb
                                orig_cursor += len(part)
                                break
                            else:
                                new_run = new_para.add_run(run_text)
                                new_run.bold = run.bold
                                new_run.italic = run.italic
                                new_run.underline = run.underline
                                new_run.font.size = run.font.size
                                new_run.font.name = run.font.name
                                new_run.font.color.rgb = run.font.color.rgb
                                orig_cursor += len(run_text)
                        else:
                            break
                    if part == "":
                        break
                    orig_cursor += 1

    return masked_doc

# Display the file content and perform masking
if uploaded_file is not None:
    if uploaded_file.type == "text/plain":
        content = uploaded_file.read().decode("utf-8")
        st.subheader("Original Text:")
        st.text(content)

        masked_content = mask_with_chatgpt(content)
        st.subheader("Masked Text:")
        st.text(masked_content)

        st.download_button(
            label="Download Masked Text",
            data=masked_content,
            file_name="masked_text.txt",
            mime="text/plain",
        )

    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        content = "\n".join([para.text for para in doc.paragraphs])
        st.subheader("Original Text:")
        st.text(content)

        masked_content = mask_with_chatgpt(content)
        masked_doc = mask_docx(doc, masked_content)

        # Save masked docx file to BytesIO
        output = BytesIO()
        masked_doc.save(output)
        output.seek(0)

        st.subheader("Masked Text:")
        st.text(masked_content)

        st.download_button(
            label="Download Masked Docx",
            data=output,
            file_name="masked_text.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
